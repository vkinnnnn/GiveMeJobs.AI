"""Document processing service implementation with LangChain integration."""

import json
import time
import io
import asyncio
from typing import Dict, Any, Optional, List
from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.prompts import PromptTemplate
from jinja2 import Template, Environment, BaseLoader
import pypdf
from docx import Document as DocxDocument

from app.core.config import get_settings
from app.core.exceptions import ProcessingException, ExternalServiceException
from app.core.logging import get_logger
from app.core.openai_client import EnhancedOpenAIClient
from app.core.dependencies import ServiceDependencies

from .models import (
    UserProfile,
    JobPosting,
    ExtractedRequirements,
    GeneratedDocument,
    DocumentFormat,
    TemplateType,
    ProcessedDocument
)

logger = get_logger(__name__)


class DocumentProcessingService:
    """Service for AI-powered document processing and generation with LangChain."""
    
    def __init__(self, dependencies: ServiceDependencies):
        self.deps = dependencies
        self.settings = dependencies.settings
        self.openai_client = dependencies.openai
        self.redis = dependencies.redis
        self.logger = dependencies.logger
        
        # Initialize LangChain components
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize LangChain ChatOpenAI
        self.langchain_llm = ChatOpenAI(
            model=self.settings.ai.openai_model,
            temperature=self.settings.ai.openai_temperature,
            max_tokens=self.settings.ai.openai_max_tokens,
            openai_api_key=self.settings.ai.openai_api_key,
            timeout=self.settings.ai.openai_timeout,
            max_retries=self.settings.ai.openai_max_retries
        )
        
        # Initialize Jinja2 environment for templates
        self.jinja_env = Environment(loader=BaseLoader())
        
        # Cache for processed documents
        self.document_cache_ttl = 3600  # 1 hour
    
    async def generate_resume(
        self,
        user_profile: UserProfile,
        job_posting: Optional[JobPosting] = None,
        template_id: TemplateType = TemplateType.PROFESSIONAL,
        custom_instructions: Optional[str] = None
    ) -> GeneratedDocument:
        """Generate AI-powered resume with LangChain and context awareness."""
        start_time = time.time()
        
        try:
            self.logger.info(
                "Starting resume generation with LangChain",
                user_id=user_profile.id,
                job_id=job_posting.id if job_posting else None,
                template=template_id
            )
            
            # Check cache first
            cache_key = self._get_resume_cache_key(user_profile.id, job_posting, template_id)
            cached_resume = await self._get_cached_document(cache_key)
            if cached_resume:
                self.logger.info("Returning cached resume", cache_key=cache_key)
                return cached_resume
            
            # Extract job requirements if job posting is provided
            job_requirements = None
            if job_posting:
                job_requirements = await self._extract_job_requirements_with_langchain(
                    job_posting.description
                )
            
            # Generate resume using LangChain
            resume_content = await self._generate_resume_with_langchain(
                user_profile, job_requirements, template_id, custom_instructions
            )
            
            # Apply template formatting
            formatted_content = await self._apply_template_formatting_async(
                resume_content, template_id, user_profile
            )
            
            generation_time = time.time() - start_time
            word_count = len(formatted_content.split())
            
            # Create document
            generated_document = GeneratedDocument(
                content=formatted_content,
                format=DocumentFormat.TXT,
                metadata={
                    "job_id": job_posting.id if job_posting else None,
                    "template_id": template_id,
                    "custom_instructions": custom_instructions is not None,
                    "langchain_model": self.settings.ai.openai_model,
                    "job_requirements_extracted": job_requirements is not None,
                    "cache_key": cache_key
                },
                generation_time=generation_time,
                word_count=word_count,
                template_used=template_id
            )
            
            # Cache the generated document
            await self._cache_document(cache_key, generated_document)
            
            self.logger.info(
                "Resume generation completed successfully",
                generation_time=generation_time,
                word_count=word_count,
                cached=True
            )
            
            return generated_document
            
        except Exception as e:
            self.logger.error("Resume generation failed", error=str(e), exc_info=True)
            raise ProcessingException(
                f"Resume generation failed: {str(e)}",
                processing_type="resume_generation",
                details={"user_id": user_profile.id}
            )
    
    async def extract_job_requirements(self, job_description: str) -> ExtractedRequirements:
        """Extract structured requirements from job description using NLP."""
        try:
            logger.info("Starting job requirement extraction")
            
            return await self._extract_job_requirements_internal(job_description)
            
        except Exception as e:
            logger.error("Job requirement extraction failed", error=str(e), exc_info=True)
            raise ProcessingException(f"Job requirement extraction failed: {str(e)}")
    
    async def _extract_job_requirements_internal(self, job_description: str) -> ExtractedRequirements:
        """Internal method for job requirement extraction."""
        prompt = f"""
        Analyze the following job description and extract structured information.
        Return the response as a valid JSON object with the following structure:
        {{
            "required_skills": ["skill1", "skill2"],
            "preferred_skills": ["skill1", "skill2"],
            "experience_years": 3,
            "education_level": "Bachelor's degree",
            "certifications": ["cert1", "cert2"],
            "responsibilities": ["responsibility1", "responsibility2"],
            "qualifications": ["qualification1", "qualification2"],
            "company_culture": ["culture1", "culture2"],
            "benefits": ["benefit1", "benefit2"],
            "employment_type": "Full-time",
            "location_requirements": "Remote/On-site/Hybrid",
            "salary_range": {{"min": 50000, "max": 80000}}
        }}
        
        Job Description:
        {job_description}
        
        Extract only the information that is explicitly mentioned or can be reasonably inferred.
        Use null for missing information.
        """
        
        try:
            response = await self.openai_client.chat_completion(
                messages=[{"role": "user", "content": prompt}],
                model=self.settings.ai.openai_model,
                temperature=0.3,  # Lower temperature for more consistent extraction
                max_tokens=1500
            )
            
            # Parse JSON response
            extracted_data = json.loads(response.choices[0].message.content)
            
            return ExtractedRequirements(**extracted_data)
            
        except json.JSONDecodeError as e:
            logger.error("Failed to parse extracted requirements JSON", error=str(e))
            # Return empty requirements if parsing fails
            return ExtractedRequirements()
        except Exception as e:
            logger.error("OpenAI API call failed", error=str(e))
            raise ExternalServiceException("OpenAI", str(e))
    
    def _build_resume_prompt(
        self,
        user_profile: UserProfile,
        job_requirements: Optional[ExtractedRequirements],
        template_id: TemplateType,
        custom_instructions: Optional[str]
    ) -> str:
        """Build context-aware prompt for resume generation."""
        
        # Base user information
        user_info = f"""
        User Profile:
        - Name: {user_profile.first_name} {user_profile.last_name}
        - Email: {user_profile.email}
        - Phone: {user_profile.phone or 'Not provided'}
        - Location: {user_profile.location or 'Not provided'}
        - Skills: {', '.join(user_profile.skills)}
        - Career Goals: {user_profile.career_goals or 'Not specified'}
        - Summary: {user_profile.summary or 'Not provided'}
        
        Experience:
        """
        
        # Add experience details
        for i, exp in enumerate(user_profile.experience, 1):
            user_info += f"""
        {i}. {exp.get('title', 'Unknown')} at {exp.get('company', 'Unknown')}
           Duration: {exp.get('duration', 'Not specified')}
           Description: {exp.get('description', 'Not provided')}
        """
        
        # Add education details
        user_info += "\nEducation:\n"
        for i, edu in enumerate(user_profile.education, 1):
            user_info += f"""
        {i}. {edu.get('degree', 'Unknown')} from {edu.get('institution', 'Unknown')}
           Year: {edu.get('year', 'Not specified')}
        """
        
        # Job requirements context
        job_context = ""
        if job_requirements:
            job_context = f"""
        
        Job Requirements to Address:
        - Required Skills: {', '.join(job_requirements.required_skills)}
        - Preferred Skills: {', '.join(job_requirements.preferred_skills)}
        - Experience Required: {job_requirements.experience_years or 'Not specified'} years
        - Education Level: {job_requirements.education_level or 'Not specified'}
        - Key Responsibilities: {', '.join(job_requirements.responsibilities)}
        """
        
        # Template-specific instructions
        template_instructions = self._get_template_instructions(template_id)
        
        # Custom instructions
        custom_context = f"\nCustom Instructions: {custom_instructions}" if custom_instructions else ""
        
        prompt = f"""
        Generate a professional resume tailored for the specific job opportunity (if provided).
        
        {user_info}
        {job_context}
        
        Template Style: {template_id.value}
        {template_instructions}
        
        Instructions:
        1. Create a compelling professional summary that highlights relevant experience
        2. Emphasize skills and experience that match the job requirements (if provided)
        3. Use action verbs and quantifiable achievements where possible
        4. Maintain professional tone and ATS-friendly formatting
        5. Keep the resume to 1-2 pages maximum
        6. Structure the resume with clear sections: Summary, Experience, Education, Skills
        7. Tailor the content to be relevant to the target role
        
        {custom_context}
        
        Generate the complete resume content in a well-structured format.
        """
        
        return prompt
    
    def _get_template_instructions(self, template_id: TemplateType) -> str:
        """Get template-specific formatting instructions."""
        templates = {
            TemplateType.MODERN: "Use a clean, modern format with clear section headers and bullet points.",
            TemplateType.CLASSIC: "Use a traditional, conservative format with standard fonts and minimal styling.",
            TemplateType.CREATIVE: "Use a more creative format with some visual elements, but keep it professional.",
            TemplateType.MINIMAL: "Use a minimalist approach with clean lines and plenty of white space.",
            TemplateType.PROFESSIONAL: "Use a standard professional format suitable for corporate environments."
        }
        return templates.get(template_id, templates[TemplateType.PROFESSIONAL])
    
    def _apply_template_formatting(self, content: str, template_id: TemplateType) -> str:
        """Apply template-specific formatting to the generated content."""
        # For now, return the content as-is
        # In a full implementation, this would apply specific formatting
        # based on the template type (fonts, spacing, layout, etc.)
        
        # Add template header comment
        formatted_content = f"<!-- Template: {template_id.value} -->\n\n{content}"
        
        return formatted_content
    
    async def process_document(
        self,
        file_content: bytes,
        file_name: str,
        format: DocumentFormat
    ) -> ProcessedDocument:
        """Process uploaded document and extract text."""
        start_time = time.time()
        
        try:
            logger.info("Starting document processing", file_name=file_name, format=format)
            
            # Save file temporarily for processing
            temp_file_path = f"/tmp/{file_name}"
            with open(temp_file_path, "wb") as f:
                f.write(file_content)
            
            # Extract text based on format
            extracted_text = ""
            metadata = {}
            
            if format == DocumentFormat.PDF:
                loader = PyPDFLoader(temp_file_path)
                documents = loader.load()
                extracted_text = "\n".join([doc.page_content for doc in documents])
                metadata["pages"] = len(documents)
                
            elif format == DocumentFormat.DOCX:
                loader = UnstructuredWordDocumentLoader(temp_file_path)
                documents = loader.load()
                extracted_text = "\n".join([doc.page_content for doc in documents])
                
            elif format == DocumentFormat.TXT:
                with open(temp_file_path, "r", encoding="utf-8") as f:
                    extracted_text = f.read()
            
            processing_time = time.time() - start_time
            
            # Clean up temporary file
            import os
            os.remove(temp_file_path)
            
            logger.info(
                "Document processing completed",
                processing_time=processing_time,
                text_length=len(extracted_text)
            )
            
            return ProcessedDocument(
                file_name=file_name,
                format=format,
                extracted_text=extracted_text,
                metadata=metadata,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error("Document processing failed", error=str(e), exc_info=True)
            raise ProcessingException(f"Document processing failed: {str(e)}")
    
    async def generate_cover_letter(
        self,
        user_profile: UserProfile,
        job_posting: JobPosting,
        custom_instructions: Optional[str] = None
    ) -> GeneratedDocument:
        """Generate AI-powered cover letter."""
        start_time = time.time()
        
        try:
            logger.info(
                "Starting cover letter generation",
                user_id=user_profile.id,
                job_id=job_posting.id
            )
            
            prompt = f"""
            Generate a professional cover letter for the following job application.
            
            User Information:
            - Name: {user_profile.first_name} {user_profile.last_name}
            - Email: {user_profile.email}
            - Skills: {', '.join(user_profile.skills)}
            - Career Goals: {user_profile.career_goals or 'Not specified'}
            
            Job Information:
            - Company: {job_posting.company}
            - Position: {job_posting.title}
            - Job Description: {job_posting.description}
            
            Instructions:
            1. Address the letter to the hiring manager
            2. Express genuine interest in the position and company
            3. Highlight relevant skills and experience that match the job requirements
            4. Show knowledge of the company and role
            5. Include a strong closing with call to action
            6. Keep it concise (3-4 paragraphs)
            7. Maintain professional tone throughout
            
            {f'Custom Instructions: {custom_instructions}' if custom_instructions else ''}
            
            Generate a complete, professional cover letter.
            """
            
            response = await self.openai_client.chat.completions.create(
                model=self.settings.openai_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=self.settings.openai_temperature,
                max_tokens=1000
            )
            
            cover_letter_content = response.choices[0].message.content
            generation_time = time.time() - start_time
            word_count = len(cover_letter_content.split())
            
            logger.info(
                "Cover letter generation completed successfully",
                generation_time=generation_time,
                word_count=word_count
            )
            
            return GeneratedDocument(
                content=cover_letter_content,
                format=DocumentFormat.TXT,
                metadata={
                    "job_id": job_posting.id,
                    "company": job_posting.company,
                    "position": job_posting.title,
                    "tokens_used": response.usage.total_tokens if response.usage else 0
                },
                generation_time=generation_time,
                word_count=word_count,
                template_used=TemplateType.PROFESSIONAL
            )
            
        except Exception as e:
            logger.error("Cover letter generation failed", error=str(e), exc_info=True)
            raise ProcessingException(f"Cover letter generation failed: {str(e)}")
    
    # New LangChain-based methods
    
    async def _generate_resume_with_langchain(
        self,
        user_profile: UserProfile,
        job_requirements: Optional[ExtractedRequirements],
        template_id: TemplateType,
        custom_instructions: Optional[str]
    ) -> str:
        """Generate resume using LangChain with structured prompts."""
        
        # Create system message for context
        system_message = SystemMessage(content="""
        You are an expert resume writer with 15+ years of experience helping professionals 
        create compelling resumes that get interviews. You understand ATS systems, 
        industry-specific requirements, and modern hiring practices.
        
        Your task is to create a professional, ATS-friendly resume that:
        1. Highlights relevant skills and experience for the target role
        2. Uses strong action verbs and quantifiable achievements
        3. Follows best practices for the specified template style
        4. Optimizes for both human readers and ATS systems
        5. Maintains professional tone and formatting
        """)
        
        # Build detailed prompt with user information
        prompt_template = PromptTemplate(
            input_variables=[
                "user_name", "user_email", "user_phone", "user_location",
                "skills", "experience", "education", "achievements",
                "career_goals", "summary", "job_requirements", "template_style",
                "custom_instructions"
            ],
            template="""
            Create a professional resume for:
            
            PERSONAL INFORMATION:
            Name: {user_name}
            Email: {user_email}
            Phone: {user_phone}
            Location: {user_location}
            
            PROFESSIONAL SUMMARY:
            {summary}
            
            CAREER GOALS:
            {career_goals}
            
            SKILLS:
            {skills}
            
            EXPERIENCE:
            {experience}
            
            EDUCATION:
            {education}
            
            ACHIEVEMENTS:
            {achievements}
            
            TARGET JOB REQUIREMENTS (if provided):
            {job_requirements}
            
            TEMPLATE STYLE:
            {template_style}
            
            CUSTOM INSTRUCTIONS:
            {custom_instructions}
            
            Generate a complete, professional resume that maximizes the candidate's 
            chances of getting an interview. Focus on relevance, impact, and readability.
            """
        )
        
        # Prepare template variables
        template_vars = {
            "user_name": f"{user_profile.first_name} {user_profile.last_name}",
            "user_email": user_profile.email,
            "user_phone": user_profile.phone or "Not provided",
            "user_location": user_profile.location or "Not provided",
            "skills": ", ".join(user_profile.skills) if user_profile.skills else "Not specified",
            "experience": self._format_experience_for_prompt(user_profile.experience),
            "education": self._format_education_for_prompt(user_profile.education),
            "achievements": ", ".join(user_profile.achievements) if user_profile.achievements else "Not specified",
            "career_goals": user_profile.career_goals or "Not specified",
            "summary": user_profile.summary or "Not provided",
            "job_requirements": self._format_job_requirements_for_prompt(job_requirements) if job_requirements else "Not provided",
            "template_style": self._get_template_description(template_id),
            "custom_instructions": custom_instructions or "None"
        }
        
        # Generate prompt
        formatted_prompt = prompt_template.format(**template_vars)
        human_message = HumanMessage(content=formatted_prompt)
        
        # Generate resume using LangChain
        response = await asyncio.to_thread(
            self.langchain_llm.invoke,
            [system_message, human_message]
        )
        
        return response.content
    
    async def _extract_job_requirements_with_langchain(
        self, job_description: str
    ) -> ExtractedRequirements:
        """Extract job requirements using LangChain with structured output."""
        
        system_message = SystemMessage(content="""
        You are an expert HR analyst specializing in job requirement extraction.
        Your task is to analyze job descriptions and extract structured information
        that can be used for resume optimization and candidate matching.
        
        Always return valid JSON with the specified structure.
        """)
        
        prompt_template = PromptTemplate(
            input_variables=["job_description"],
            template="""
            Analyze the following job description and extract structured information.
            Return ONLY a valid JSON object with this exact structure:
            
            {{
                "required_skills": ["skill1", "skill2"],
                "preferred_skills": ["skill1", "skill2"],
                "experience_years": 3,
                "education_level": "Bachelor's degree",
                "certifications": ["cert1", "cert2"],
                "responsibilities": ["responsibility1", "responsibility2"],
                "qualifications": ["qualification1", "qualification2"],
                "company_culture": ["culture1", "culture2"],
                "benefits": ["benefit1", "benefit2"],
                "employment_type": "Full-time",
                "location_requirements": "Remote/On-site/Hybrid",
                "salary_range": {{"min": 50000, "max": 80000}}
            }}
            
            Job Description:
            {job_description}
            
            Extract only information that is explicitly mentioned or can be reasonably inferred.
            Use null for missing information. Ensure the response is valid JSON.
            """
        )
        
        formatted_prompt = prompt_template.format(job_description=job_description)
        human_message = HumanMessage(content=formatted_prompt)
        
        try:
            response = await asyncio.to_thread(
                self.langchain_llm.invoke,
                [system_message, human_message]
            )
            
            # Parse JSON response
            extracted_data = json.loads(response.content)
            return ExtractedRequirements(**extracted_data)
            
        except json.JSONDecodeError as e:
            self.logger.error("Failed to parse LangChain extracted requirements JSON", error=str(e))
            return ExtractedRequirements()
        except Exception as e:
            self.logger.error("LangChain job requirement extraction failed", error=str(e))
            raise ProcessingException(f"Job requirement extraction failed: {str(e)}")
    
    async def _apply_template_formatting_async(
        self, content: str, template_id: TemplateType, user_profile: UserProfile
    ) -> str:
        """Apply template-specific formatting using Jinja2 templates."""
        
        # Define template styles
        templates = {
            TemplateType.MODERN: """
{{ content }}

---
Template: Modern Professional
Generated for: {{ user_name }}
Date: {{ generation_date }}
            """,
            TemplateType.CLASSIC: """
{{ content }}

---
Template: Classic Professional
Generated for: {{ user_name }}
Date: {{ generation_date }}
            """,
            TemplateType.CREATIVE: """
{{ content }}

---
Template: Creative Professional
Generated for: {{ user_name }}
Date: {{ generation_date }}
            """,
            TemplateType.MINIMAL: """
{{ content }}

---
Template: Minimal Professional
Generated for: {{ user_name }}
Date: {{ generation_date }}
            """,
            TemplateType.PROFESSIONAL: """
{{ content }}

---
Template: Professional
Generated for: {{ user_name }}
Date: {{ generation_date }}
            """
        }
        
        template_str = templates.get(template_id, templates[TemplateType.PROFESSIONAL])
        template = self.jinja_env.from_string(template_str)
        
        formatted_content = template.render(
            content=content,
            user_name=f"{user_profile.first_name} {user_profile.last_name}",
            generation_date=time.strftime("%Y-%m-%d %H:%M:%S")
        )
        
        return formatted_content
    
    def _format_experience_for_prompt(self, experience: List[Dict[str, Any]]) -> str:
        """Format experience list for prompt."""
        if not experience:
            return "No experience provided"
        
        formatted = []
        for i, exp in enumerate(experience, 1):
            formatted.append(f"""
            {i}. {exp.get('title', 'Unknown Position')} at {exp.get('company', 'Unknown Company')}
               Duration: {exp.get('duration', 'Not specified')}
               Description: {exp.get('description', 'Not provided')}
               Achievements: {exp.get('achievements', 'Not specified')}
            """)
        
        return "\n".join(formatted)
    
    def _format_education_for_prompt(self, education: List[Dict[str, Any]]) -> str:
        """Format education list for prompt."""
        if not education:
            return "No education provided"
        
        formatted = []
        for i, edu in enumerate(education, 1):
            formatted.append(f"""
            {i}. {edu.get('degree', 'Unknown Degree')} from {edu.get('institution', 'Unknown Institution')}
               Year: {edu.get('year', 'Not specified')}
               GPA: {edu.get('gpa', 'Not provided')}
               Relevant Coursework: {edu.get('coursework', 'Not specified')}
            """)
        
        return "\n".join(formatted)
    
    def _format_job_requirements_for_prompt(self, requirements: ExtractedRequirements) -> str:
        """Format job requirements for prompt."""
        return f"""
        Required Skills: {', '.join(requirements.required_skills)}
        Preferred Skills: {', '.join(requirements.preferred_skills)}
        Experience Required: {requirements.experience_years} years
        Education Level: {requirements.education_level}
        Key Responsibilities: {', '.join(requirements.responsibilities)}
        Employment Type: {requirements.employment_type}
        Location: {requirements.location_requirements}
        """
    
    # Caching methods
    
    def _get_resume_cache_key(
        self, user_id: str, job_posting: Optional[JobPosting], template_id: TemplateType
    ) -> str:
        """Generate cache key for resume."""
        job_id = job_posting.id if job_posting else "no_job"
        return f"resume:{user_id}:{job_id}:{template_id.value}"
    
    async def _get_cached_document(self, cache_key: str) -> Optional[GeneratedDocument]:
        """Get document from cache."""
        try:
            cached_data = await self.redis.get(cache_key)
            if cached_data:
                data = json.loads(cached_data)
                return GeneratedDocument(**data)
        except Exception as e:
            self.logger.warning("Failed to get cached document", cache_key=cache_key, error=str(e))
        return None
    
    async def _cache_document(self, cache_key: str, document: GeneratedDocument) -> None:
        """Cache generated document."""
        try:
            data = document.model_dump_json()
            await self.redis.setex(cache_key, self.document_cache_ttl, data)
        except Exception as e:
            self.logger.warning("Failed to cache document", cache_key=cache_key, error=str(e))
    
    # Document processing with LangChain
    
    async def process_document_with_langchain(
        self, file_content: bytes, file_name: str, format: DocumentFormat
    ) -> ProcessedDocument:
        """Process document using LangChain loaders with enhanced text extraction."""
        start_time = time.time()
        
        try:
            self.logger.info("Processing document with LangChain", file_name=file_name, format=format)
            
            # Create temporary file
            temp_file_path = Path(f"/tmp/{file_name}")
            temp_file_path.parent.mkdir(exist_ok=True)
            
            with open(temp_file_path, "wb") as f:
                f.write(file_content)
            
            extracted_text = ""
            metadata = {}
            
            if format == DocumentFormat.PDF:
                # Use PyPDF for better text extraction
                pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
                pages = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    pages.append(page_text)
                
                extracted_text = "\n\n".join(pages)
                metadata = {
                    "pages": len(pages),
                    "pdf_info": pdf_reader.metadata,
                    "extraction_method": "pypdf"
                }
                
            elif format == DocumentFormat.DOCX:
                # Use python-docx for better structure preservation
                doc = DocxDocument(io.BytesIO(file_content))
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                
                extracted_text = "\n".join(paragraphs)
                metadata = {
                    "paragraphs": len(paragraphs),
                    "extraction_method": "python-docx"
                }
                
            elif format == DocumentFormat.TXT:
                extracted_text = file_content.decode('utf-8', errors='ignore')
                metadata = {"extraction_method": "direct"}
            
            # Clean up temporary file
            if temp_file_path.exists():
                temp_file_path.unlink()
            
            # Use LangChain text splitter for better chunking
            if len(extracted_text) > 2000:
                chunks = self.text_splitter.split_text(extracted_text)
                metadata["chunks"] = len(chunks)
                metadata["chunked"] = True
            else:
                metadata["chunked"] = False
            
            processing_time = time.time() - start_time
            
            self.logger.info(
                "Document processing with LangChain completed",
                processing_time=processing_time,
                text_length=len(extracted_text),
                chunks=metadata.get("chunks", 1)
            )
            
            return ProcessedDocument(
                file_name=file_name,
                format=format,
                extracted_text=extracted_text,
                metadata=metadata,
                processing_time=processing_time
            )
            
        except Exception as e:
            self.logger.error("LangChain document processing failed", error=str(e), exc_info=True)
            raise ProcessingException(
                f"Document processing failed: {str(e)}",
                processing_type="document_processing"
            )